import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from PIL import Image, ImageDraw, ImageFont

def set_cell_background(cell, fill_hex):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner padding for table cells."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    """
    Set cell borders: top, bottom, left, right.
    Format: top={"sz": 4, "val": "single", "color": "CCCCCC"}
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name, border_props in kwargs.items():
        if border_props is None:
            node = OxmlElement(f'w:{border_name}')
            node.set(qn('w:val'), 'none')
            tcBorders.append(node)
        else:
            node = OxmlElement(f'w:{border_name}')
            node.set(qn('w:val'), border_props.get('val', 'single'))
            node.set(qn('w:sz'), str(border_props.get('sz', 4)))
            node.set(qn('w:space'), '0')
            node.set(qn('w:color'), border_props.get('color', 'CCCCCC'))
            tcBorders.append(node)
    tcPr.append(tcBorders)

def generate_logo():
    """Generate cover logo image matching reference PDF style."""
    os.makedirs("assets", exist_ok=True)
    logo_path = "assets/logo.png"
    img = Image.new("RGBA", (300, 300), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    
    # Draw geometric flower star matching reference PDF logo
    center = (150, 150)
    # Outer hexagon/star petals in purple #6B21A8 / #581C87
    # Inner elements in gold/orange #F59E0B
    import math
    for i in range(6):
        angle = i * (2 * math.pi / 6)
        r_outer = 120
        r_inner = 50
        pts = [
            (center[0] + r_outer * math.cos(angle), center[1] + r_outer * math.sin(angle)),
            (center[0] + r_inner * math.cos(angle + math.pi/6), center[1] + r_inner * math.sin(angle + math.pi/6)),
            (center[0] + r_outer * 0.7 * math.cos(angle + math.pi/3), center[1] + r_outer * 0.7 * math.sin(angle + math.pi/3)),
        ]
        draw.polygon(pts, fill="#6B21A8")

    # Draw inner yellow ring and orange circle
    draw.ellipse([100, 100, 200, 200], fill="#FEF08A", outline="#EAB308", width=4)
    draw.ellipse([120, 120, 180, 180], fill="#F97316")
    draw.ellipse([135, 135, 165, 165], fill="#EA580C")
    
    img.save(logo_path)
    return logo_path

def add_code_block(doc, code_text):
    """Add a shaded code block table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F4F4F5")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_border(cell, 
                    top={"sz": 4, "val": "single", "color": "E4E4E7"},
                    bottom={"sz": 4, "val": "single", "color": "E4E4E7"},
                    left={"sz": 12, "val": "single", "color": "10B981"}, # emerald left accent line
                    right={"sz": 4, "val": "single", "color": "E4E4E7"})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    
    lines = code_text.strip().split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x27, 0x27, 0x2A)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_framed_image(doc, img_path, width_inches=6.0):
    """Add an image placed neatly inside a table frame matching PDF layout."""
    if not os.path.exists(img_path):
        p = doc.add_paragraph(f"[Gambar {img_path} tidak ditemukan]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
        
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
    set_cell_border(cell,
                    top={"sz": 4, "val": "single", "color": "CBD5E1"},
                    bottom={"sz": 4, "val": "single", "color": "CBD5E1"},
                    left={"sz": 4, "val": "single", "color": "CBD5E1"},
                    right={"sz": 4, "val": "single", "color": "CBD5E1"})
                    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(10)

def create_document():
    doc = Document()

    # Set page margins to 1 inch (72 pt / 2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # -------------------------------------------------------------
    # COVER PAGE
    # -------------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(4)
    r = p_title.add_run("LAPORAN PRAKTIKUM PEMROGRAMAN WEB")
    r.font.bold = True
    r.font.size = Pt(16)

    p_sub1 = doc.add_paragraph()
    p_sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub1.paragraph_format.space_after = Pt(4)
    r = p_sub1.add_run("FULL-STACK CRUD LENGKAP")
    r.font.bold = True
    r.font.size = Pt(16)

    p_sub2 = doc.add_paragraph()
    p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub2.paragraph_format.space_after = Pt(36)
    r = p_sub2.add_run("(CREATE, READ, UPDATE, DELETE)")
    r.font.bold = True
    r.font.size = Pt(14)

    # Logo Graphic
    logo_file = generate_logo()
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(48)
    p_logo.add_run().add_picture(logo_file, width=Inches(2.5))

    # Metadata Block
    meta_items = [
        ("Nama", ": Arya Leksana Puntadewa"),
        ("NIM", ": 24.11.6109"),
        ("Kelas", ": 24 S1 IF03"),
        ("Judul Tugas", ": Tugas 11 — Full-Stack Todo App (Nordic Minimalist Light & Emerald)")
    ]

    table_meta = doc.add_table(rows=4, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, val) in enumerate(meta_items):
        cell_lbl = table_meta.cell(i, 0)
        cell_val = table_meta.cell(i, 1)
        
        # remove borders
        set_cell_border(cell_lbl, top=None, bottom=None, left=None, right=None)
        set_cell_border(cell_val, top=None, bottom=None, left=None, right=None)
        
        p_l = cell_lbl.paragraphs[0]
        p_l.paragraph_format.space_after = Pt(4)
        p_l.paragraph_format.line_spacing = 1.15
        r_l = p_l.add_run(label)
        r_l.font.bold = True
        r_l.font.size = Pt(12)
        
        p_v = cell_val.paragraphs[0]
        p_v.paragraph_format.space_after = Pt(4)
        p_v.paragraph_format.line_spacing = 1.15
        r_v = p_v.add_run(val)
        r_v.font.size = Pt(12)
        
    doc.add_page_break()

    # Helper function for section headings
    def add_sec_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(14)

    def add_sub_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(12)

    def add_subsub_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(12)

    def add_body_p(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.size = Pt(12)
        return p

    # -------------------------------------------------------------
    # A. HASIL DAN PEMBAHASAN
    # -------------------------------------------------------------
    add_sec_heading("A. Hasil dan Pembahasan")
    add_body_p("Bagian ini menunjukkan hasil implementasi aplikasi Full-Stack CRUD Todo App yang menghubungkan backend NestJS + MySQL/Data Persistence dengan frontend React (Vite + Axios) dengan tema desain Nordic Minimalist Light & Emerald.")

    # 1. Form Add Todo
    add_sub_heading("1. Form Add Todo")
    add_body_p("Form modal untuk menambah todo baru. Saat mode tambah dibuka, semua field input (Judul Tugas, Deskripsi, Prioritas, Tanggal Tenggat) berada dalam kondisi kosong/default, dan header modal menampilkan indikator '+ Tambah Todo Baru'. Validasi dilakukan baik di frontend maupun backend.")
    add_framed_image(doc, "screenshots/01_add_todo_form.png")

    # 2. Form Edit Todo (Prepopulate)
    add_sub_heading("2. Form Edit Todo (Prepopulate)")
    add_body_p("Form modal yang sama digunakan untuk mode edit todo. Saat tombol \"Edit\" pada salah satu kartu todo diklik, state editingTodo terisi dan form modal otomatis terisi (prepopulate) dengan data todo terkait (Judul, Deskripsi, Prioritas, Status, dan Tanggal Tenggat), serta badge modal berubah menjadi mode Edit.")
    add_framed_image(doc, "screenshots/02_edit_todo_form.png")

    # 3. Daftar Produk / Todo
    add_sub_heading("3. Daftar Todo (Grid View)")
    add_body_p("Seluruh todo ditampilkan dalam bentuk grid kartu (card layout) yang responsif dan intuitif. Masing-masing kartu menampilkan judul tugas, deskripsi, tanggal tenggat, badge prioritas (High - Merah, Medium - Kuning, Low - Biru), status penyelesaian (Selesai / Belum Selesai), serta tombol aksi Edit dan Hapus.")
    add_framed_image(doc, "screenshots/03_todo_list.png")

    # 4. Toggle Availability / Complete Status
    add_sub_heading("4. Quick Toggle Complete Status")
    add_body_p("Tombol/badge status 'Selesai' atau 'Belum Selesai' pada setiap kartu todo dapat diklik langsung untuk mengubah status penyelesaian tugas (mengirim request PUT ke /api/v1/todos/:id), secara cepat tanpa perlu membuka form modal edit lengkap.")
    add_framed_image(doc, "screenshots/04_toggle_complete.png")

    # 5. Search & Filter
    add_sub_heading("5. Search & Filter Data")
    add_body_p("Fitur pencarian menyaring todo secara real-time berdasarkan judul dan deskripsi. Selain itu, terdapat tab filter status ('Semua', 'Aktif', 'Selesai') dan dropdown filter prioritas ('Semua Prioritas', 'High', 'Medium', 'Low') yang memudahkan manajemen tugas.")
    add_framed_image(doc, "screenshots/05_search_filter.png")

    # 6. Toast Notification & Loading Indicator
    add_sub_heading("6. Toast Notification & Indikator Loading")
    add_body_p("Setiap kali aksi CRUD berhasil atau gagal dijalankan, sistem secara otomatis memunculkan Toast Notification notifikasi melayang di pojok layar. Selain itu, saat data sedang diambil dari server backend NestJS, indikator loading ditampilkan untuk memberikan kepastian antarmuka kepada pengguna.")
    add_framed_image(doc, "screenshots/06_toast_loading.png")

    # 7. Error Handling
    add_sub_heading("7. Error Handling")
    add_body_p("Ketika backend NestJS tidak dapat dihubungi (misalnya server mati atau masalah koneksi jaringan), aplikasi menangani exception secara graceful dengan menampilkan banner error pesan yang jelas beserta tombol 'Coba lagi' untuk melakukan refetch data.")
    add_framed_image(doc, "screenshots/07_error_handling.png")

    # 8. Log Console — Axios Interceptor
    add_sub_heading("8. Log Console — Axios Interceptor")
    add_body_p("Axios Request Interceptor mencatat setiap HTTP method (GET, POST, PUT, DELETE) dan endpoint URL yang dikirim ke server. Sedangkan Axios Response Interceptor mencatat status response (HTTP 200 OK, 201 Created) serta menangani format error secara terpusat di console browser.")
    add_framed_image(doc, "screenshots/08_console_interceptor.png")

    # -------------------------------------------------------------
    # B. PENJELASAN KODE
    # -------------------------------------------------------------
    add_sec_heading("B. Penjelasan Kode")
    
    add_sub_heading("1. Flow CRUD (Create, Read, Update, Delete)")
    
    add_subsub_heading("Create (POST)")
    add_body_p("Saat form modal dalam mode tambah, pengiriman form memanggil handler handleFormSuccess, yang mengirim HTTP request POST ke /api/v1/todos membawa payload data todo baru. Setelah backend NestJS merespons berhasil, fungsi fetchTodos() dipanggil ulang agar daftar todo di layar sinkron secara penuh dengan database.")
    add_code_block(doc, 
"""// Controller NestJS: todo.controller.ts
@Post()
create(@Body() createTodoDto: CreateTodoDto) {
  return this.todoService.create(createTodoDto);
}

// Frontend React: TodoForm.jsx
const handleSubmit = async (e) => {
  e.preventDefault();
  await api.post('/api/v1/todos', formData);
  onSuccess('✨ Todo baru berhasil ditambahkan!');
};""")

    add_subsub_heading("Read (GET)")
    add_body_p("Fungsi fetchTodos dipanggil saat komponen pertama kali dimuat (di dalam useEffect hook), serta dipanggil ulang setelah setiap operasi Create/Update/Delete selesai. Mendukung query parameter filter completed (?completed=true atau ?completed=false) yang diproses langsung oleh NestJS controller.")
    add_code_block(doc,
"""// Controller NestJS: todo.controller.ts
@Get()
findAll(@Query('completed') completed?: string) {
  let completedFilter: boolean | undefined;
  if (completed === 'true') completedFilter = true;
  if (completed === 'false') completedFilter = false;
  return this.todoService.findAll(completedFilter);
}

// Frontend React: TodoList.jsx
const fetchTodos = useCallback(async () => {
  let url = '/api/v1/todos';
  if (filterStatus === 'active') url += '?completed=false';
  if (filterStatus === 'completed') url += '?completed=true';

  const response = await api.get(url);
  setTodos(response.data);
}, [filterStatus]);""")

    add_subsub_heading("Update (PUT) — Edit Form")
    add_body_p("Saat tombol Edit pada kartu diklik, state editingTodo diisi dengan objek todo tersebut. Modal form terisi otomatis via prepopulate. Submit form pada mode ini mengirim request PUT ke /api/v1/todos/:id dan melakukan refetch data.")
    add_code_block(doc,
"""// Frontend React: TodoForm.jsx
if (isEdit) {
  await api.put(`/api/v1/todos/${todo.id}`, formData);
  onSuccess('🎉 Todo berhasil diperbarui!');
}""")

    add_subsub_heading("Update (PUT) — Quick Toggle Complete")
    add_body_p("Tombol status pada kartu todo memanggil handleToggleComplete, mengirimkan request PUT berisi field { completed: !todo.completed } — hanya mengubah status penyelesaian tanpa harus membuka modal form edit.")
    add_code_block(doc,
"""// Frontend React: TodoList.jsx
const handleToggleComplete = async (todo) => {
  const updatedStatus = !todo.completed;
  await api.put(`/api/v1/todos/${todo.id}`, { completed: updatedStatus });
  showToast(updatedStatus ? '✅ Status diubah menjadi Selesai!' : '🔄 Status diubah menjadi Belum Selesai');
  fetchTodos(); // Refetch
};""")

    add_subsub_heading("Delete (DELETE)")
    add_body_p("Tombol Hapus membuka modal konfirmasi kustom (ConfirmModal). Jika pengguna mengonfirmasi, aplikasi mengirim HTTP DELETE ke /api/v1/todos/:id dan menyegarkan tampilan.")
    add_code_block(doc,
"""// Controller NestJS: todo.controller.ts
@Delete(':id')
remove(@Param('id', ParseIntPipe) id: number) {
  return this.todoService.remove(id);
}

// Frontend React: TodoList.jsx
const handleConfirmDelete = async () => {
  await api.delete(`/api/v1/todos/${deleteConfirmId}`);
  showToast('🗑️ Todo berhasil dihapus!');
  setDeleteConfirmId(null);
  fetchTodos();
};""")

    add_sub_heading("2. Cara Kerja Refetch Pattern")
    add_body_p("Alih-alih memperbarui state lokal secara manual di frontend (optimistic update), setiap handler aksi CRUD memanggil ulang fungsi fetchTodos() setelah request API selesai dengan sukses. Pendekatan ini memastikan tampilan UI selalu mencerminkan kondisi data yang sebenarnya tersimpan di backend NestJS dan database MySQL, meminimalkan risiko ketidaksinkronan UI akibat validasi atau logika server-side.")
    add_code_block(doc,
"""// Refetch Pattern
await api.put(`/api/v1/todos/${id}`, updatedData);
await fetchTodos(); // Ambil ulang data paling presisi dari server""")

    add_sub_heading("3. Reuse Komponen Form & Modal")
    add_body_p("Komponen TodoForm dipakai secara efisien untuk mode tambah dan edit sekaligus. Perbedaannya terletak pada prop todo: jika null maka mode tambah (form kosong), dan jika berisi objek todo maka mode edit (form terisi otomatis melalui hook useEffect). Komponen Modal dan ConfirmModal juga dibuat terpisah sehingga reusable untuk berbagai kebutuhan UI dialog.")

    # -------------------------------------------------------------
    # C. KESIMPULAN
    # -------------------------------------------------------------
    add_sec_heading("C. Kesimpulan")
    
    kesimpulan_list = [
        "1. Pola refetch data setelah setiap operasi CRUD (Create, Read, Update, Delete) memastikan tampilan frontend React selalu 100% sinkron dengan data yang sebenarnya tersimpan di database NestJS, tanpa perlu mengelola state manual yang rawan bug.",
        "2. Reuse satu komponen form modal (TodoForm & Modal) untuk mode tambah dan edit (menggunakan controlled component serta prepopulate via useEffect) membuat kode jauh lebih bersih, ringkas, dan maintainable dibanding membuat dua form terpisah.",
        "3. Penanganan state loading, Axios Interceptor logging terpusat, Toast Notification, serta Error Handling eksplisit di seluruh alur aplikasi menciptakan pengalaman pengguna (User Experience) yang sangat responsif, informatif, dan profesional."
    ]

    for item in kesimpulan_list:
        add_body_p(item)

    # -------------------------------------------------------------
    # D. LAMPIRAN
    # -------------------------------------------------------------
    add_sec_heading("D. Lampiran")
    
    add_sub_heading("1. Link GitHub Repository")
    p_link = doc.add_paragraph()
    p_link.paragraph_format.space_after = Pt(8)
    r_link = p_link.add_run("https://github.com/Vincrescent/11-todo")
    r_link.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8) # Blue underline link
    r_link.font.underline = True

    add_sub_heading("2. Struktur Folder Proyek")
    
    add_body_p("Backend — todo-api")
    add_code_block(doc,
"""src/
├── todo/
│   ├── dto/
│   │   ├── create-todo.dto.ts
│   │   └── update-todo.dto.ts
│   ├── entities/
│   │   └── todo.entity.ts
│   ├── todo.controller.ts
│   ├── todo.service.ts
│   └── todo.module.ts
├── app.module.ts
└── main.ts""")

    add_body_p("Frontend — todo-frontend")
    add_code_block(doc,
"""src/
├── api/
│   └── axiosInstance.js
├── components/
│   ├── ConfirmModal.jsx
│   ├── Modal.jsx
│   ├── Toast.jsx
│   ├── TodoCard.jsx
│   └── TodoForm.jsx
├── pages/
│   └── TodoList.jsx
├── App.jsx
├── index.css
└── main.jsx""")

    output_filename = "Laporan_Praktikum_11_FullStack_CRUD_Todo.docx"
    doc.save(output_filename)
    print(f"Report document created successfully: {output_filename}")

if __name__ == "__main__":
    create_document()
